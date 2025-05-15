import asyncio
import io

# import json # No longer needed for json.loads
import logging
import re
from typing import Any
from typing import Protocol
from typing import cast
from uuid import uuid4

from docx import Document
from docxtpl import DocxTemplate

from app.models.report_models import (  # Keep for potential future use or type hinting if context becomes more specific
    ReportContext,
)

# Configure module logger
logger = logging.getLogger(__name__)


class DocBuilderError(Exception):
    """Base exception for document building errors"""


# Define a Protocol for paragraph objects with clear method
class ParagraphProtocol(Protocol):
    """Protocol defining the interface for paragraph objects."""

    def clear(self) -> None: ...
    def add_run(self, text: str) -> Any: ...
    @property
    def text(self) -> str: ...
    @property
    def style(self) -> Any: ...


BOLD_RE = re.compile(r"\\*\\*(.+?)\\*\\*")

# Maps the DOCX template tags (keys) to the expected keys in the context dictionary (values).
TEMPLATE_TAG_TO_CONTEXT_KEY_MAPPING: dict[str, str] = {
    "CLIENT": "client",
    "CLIENTADDRESS1": "client_address1",
    "CLIENTADDRESS2": "client_address2",
    "DATE": "date",
    "VSRIF": "vs_rif",
    "RIFBROKER": "rif_broker",
    "POLIZZA": "polizza",
    "NSRIF": "ns_rif",
    "ASSICURATO": "assicurato",
    "INDIRIZZOASSICURATO1": "indirizzo_ass1",
    "INDIRIZZOASSICURATO2": "indirizzo_ass2",
    "LUOGO": "luogo",
    "DATADANNO": "data_danno",
    "CAUSE": "cause",
    "DATAINCARICO": "data_incarico",
    "MERCE": "merce",
    "PESOMERCE": "peso_merce",
    "VALOREMERCE": "valore_merce",
    "DATAINTERVENTO": "data_intervento",
    "ALLEGATI": "allegati",
    # Tags that will be replaced by multi-paragraph content later
    "DINAMICA_EVENTI": "{{DINAMICA_EVENTI}}",
    "ACCERTAMENTI": "{{ACCERTAMENTI}}",
    "QUANTIFICAZIONE": "{{QUANTIFICAZIONE}}",
    "COMMENTO": "{{COMMENTO}}",
}

# Maps the section placeholder tags found in the template (keys)
# to the expected keys in the context dictionary that hold their content (values).
SECTION_PLACEHOLDER_TO_CONTEXT_KEY_MAPPING: dict[str, str] = {
    "{{DINAMICA_EVENTI}}": "dinamica_eventi",
    "{{ACCERTAMENTI}}": "accertamenti",
    "{{QUANTIFICAZIONE}}": "quantificazione",
    "{{COMMENTO}}": "commento",
}


def _add_markdown(par: ParagraphProtocol, txt: str) -> None:
    """Add markdown-style formatting to a paragraph."""
    # Ensure txt is a string before processing, handles None or other types gracefully
    processed_txt = str(txt) if txt is not None else ""
    try:
        pos = 0
        for m in BOLD_RE.finditer(processed_txt):
            if m.start() > pos:
                par.add_run(processed_txt[pos : m.start()])
            par.add_run(m.group(1)).bold = True
            pos = m.end()
        if pos < len(processed_txt):
            par.add_run(processed_txt[pos:])
    except Exception as e:
        logger.error("Failed to add markdown formatting: %s", str(e), exc_info=True)
        raise DocBuilderError("Failed to apply text formatting") from e


async def inject(template_path: str, context: ReportContext) -> bytes:
    """Inject ReportContext content into the document template asynchronously."""

    def _sync_doc_generation(template_path_str: str, report_context: ReportContext) -> bytes:
        request_id = str(uuid4())
        logger.info("[%s] Starting document generation with template: %s", request_id, template_path_str)
        try:
            tpl = DocxTemplate(template_path_str)
            logger.debug("[%s] Successfully loaded template", request_id)
            mapping_data = {
                tpl_tag: getattr(report_context, ctx_key, None)
                for tpl_tag, ctx_key in TEMPLATE_TAG_TO_CONTEXT_KEY_MAPPING.items()
                if ctx_key not in SECTION_PLACEHOLDER_TO_CONTEXT_KEY_MAPPING.values()
            }
            allegati_content = getattr(report_context, "allegati", [])
            if isinstance(allegati_content, list):
                mapping_data["ALLEGATI"] = "\n".join(str(item) for item in allegati_content if item is not None)
            elif allegati_content is not None:
                mapping_data["ALLEGATI"] = str(allegati_content)
            else:
                mapping_data["ALLEGATI"] = ""
            for _tag, _key in SECTION_PLACEHOLDER_TO_CONTEXT_KEY_MAPPING.items():
                mapping_data[_tag.removeprefix("{{").removesuffix("}}")] = _tag
            tpl.render(mapping_data)
            logger.debug("[%s] Successfully rendered template with mapping_data", request_id)
            bio = io.BytesIO()
            tpl.save(bio)
            bio.seek(0)
            doc = Document(bio)
            section_content_map = {
                placeholder: getattr(report_context, ctx_key, "")
                for placeholder, ctx_key in SECTION_PLACEHOLDER_TO_CONTEXT_KEY_MAPPING.items()
            }
            for p_raw in doc.paragraphs:
                p = cast(ParagraphProtocol, p_raw)  # Cast to our Protocol type
                current_text = p.text
                for tag, _content_key_in_map in SECTION_PLACEHOLDER_TO_CONTEXT_KEY_MAPPING.items():
                    if tag in current_text:
                        content_string = section_content_map.get(tag, "")
                        style = p.style
                        p.clear()  # p is now properly typed
                        paragraphs_to_insert = [
                            t.strip()
                            for t in (str(content_string) if content_string is not None else "").split("\n\n")
                            if t.strip()
                        ]
                        if not paragraphs_to_insert:
                            p.add_run("")
                        else:
                            for idx, para_text in enumerate(paragraphs_to_insert):
                                # Cast the new paragraph to our Protocol type
                                target_par = p if idx == 0 else cast(ParagraphProtocol, doc.add_paragraph(style=style))
                                if idx > 0:
                                    target_par.clear()
                                _add_markdown(target_par, para_text)
                        logger.debug(
                            "[%s] Processed section for placeholder: %s",
                            request_id,
                            tag,
                        )
                        break
            out = io.BytesIO()
            doc.save(out)
            out.seek(0)
            result = out.read()
            logger.info(
                "[%s] Successfully generated document, size: %d bytes",
                request_id,
                len(result),
            )
            return result
        except DocBuilderError:
            raise
        except Exception as e:
            logger.exception("[%s] Unexpected error in document generation", request_id)
            raise DocBuilderError("Unexpected error in document generation") from e

    return await asyncio.to_thread(_sync_doc_generation, template_path, context)
