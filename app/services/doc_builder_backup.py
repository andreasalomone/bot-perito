import json, io, re
from docxtpl import DocxTemplate
from docx import Document

# ---------------------------------------------------------------
# Regex per **bold** in stile Markdown
# ---------------------------------------------------------------
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_markdown_line(par, line: str) -> None:
    """
    Aggiunge la stringa `line` in `par`, trasformando i segmenti racchiusi
    da **asterischi** in run grassetto.
    """
    pos = 0
    for m in BOLD_RE.finditer(line):
        # testo normale prima del blocco in grassetto
        if m.start() > pos:
            par.add_run(line[pos:m.start()])
        # blocco in grassetto
        par.add_run(m.group(1)).bold = True
        pos = m.end()

    # eventuale coda dopo l’ultimo blocco bold
    if pos < len(line):
        par.add_run(line[pos:])


# ---------------------------------------------------------------
# Funzione principale di merge
# ---------------------------------------------------------------
def inject(template_path: str, json_payload: str) -> bytes:
    """
    1. Compila i placeholder semplici tramite DocxTpl.
    2. Riapre il documento con python-docx e sostituisce {{BODY}}
       con testo multi-paragrafo, gestendo il grassetto Markdown.
    """
    # ---------- step 1: popolamento placeholder -----------------
    tpl = DocxTemplate(template_path)
    ctx = json.loads(json_payload)

    # Corpo testuale tenuto da parte per lo step 2
    body_raw = ctx.pop("body", "")

    # Mappatura chiavi JSON → placeholder DOCX (in maiuscolo)
    tpl.render({
        "CLIENT":         ctx.get("client", ""),
        "CLIENTADDRESS1": ctx.get("client_address1", ""),
        "CLIENTADDRESS2": ctx.get("client_address2", ""),
        "DATE":           ctx.get("date", ""),
        "VSRIF":          ctx.get("vs_rif", ""),
        "RIFBROKER":      ctx.get("rif_broker", ""),
        "POLIZZA":        ctx.get("polizza", ""),
        "NSRIF":          ctx.get("ns_rif", ""),
        "SUBJECT":        ctx.get("subject", ""),
        # BODY rimane come placeholder per la sostituzione successiva
        "BODY": "{{BODY}}",
    })

    # Salva in buffer memoria
    bio = io.BytesIO()
    tpl.save(bio)
    bio.seek(0)

    # ---------- step 2: sostituzione BODY -----------------------
    doc = Document(bio)
    marker = "{{BODY}}"

    for p in doc.paragraphs:
        if marker in p.text:
            base_style = p.style      # conserva lo stile paragrafo
            p.clear()                 # rimuove il placeholder

            # Split in blocchi separati da righe vuote
            blocks = [b.strip() for b in body_raw.split("\n\n") if b.strip()]
            for i, raw in enumerate(blocks):
                target_par = p if i == 0 else doc.add_paragraph(style=base_style)
                _add_markdown_line(target_par, raw)

                # ‑-- nuovo: paragrafo vuoto di separazione
                if i < len(blocks) - 1:
                    doc.add_paragraph(style=base_style)

            break  # BODY sostituito ⇒ uscita dal loop

    # ---------- output -----------------------------------------
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()