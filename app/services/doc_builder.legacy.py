import asyncio
import io

# import json # No longer needed for json.loads
import logging
from typing import Any
from typing import Protocol
from typing import cast
from uuid import uuid4

from docx import Document
from docxtpl import DocxTemplate

from app.models.report_models import ReportContext  # Keep for potential future use or type hinting if context becomes more specific

# Configure module logger
logger = logging.getLogger(__name__)


class DocBuilderError(Exception):
    """Base exception for document building errors"""


# Define a Protocol for paragraph objects with clear method
class ParagraphProtocol(Protocol):
    """Protocol defining the interface for paragraph objects."""

    def clear(self) -> None: ...
    def add_run(self, text: str) -> Any: ...
    @property
    def text(self) -> str: ...
    @property
    def style(self) -> Any: ...
    def insert_paragraph_after(self, text: str | None = ..., style: Any | None = ...) -> Any: ...


# Maps the DOCX template tags (keys) to the expected keys in the context dictionary (values).
TEMPLATE_TAG_TO_CONTEXT_KEY_MAPPING: dict[str, str] = {
    "CLIENT": "client",
    "CLIENTADDRESS1": "client_address1",
    "CLIENTADDRESS2": "client_address2",
    "DATE": "date",
    "VSRIF": "vs_rif",
    "RIFBROKER": "rif_broker",
    "POLIZZA": "polizza",
    "NSRIF": "ns_rif",
    "ASSICURATO": "assicurato",
    "INDIRIZZOASSICURATO1": "indirizzo_ass1",
    "INDIRIZZOASSICURATO2": "indirizzo_ass2",
    "LUOGO": "luogo",
    "DATADANNO": "data_danno",
    "CAUSE": "cause",
    "DATAINCARICO": "data_incarico",
    "MERCE": "merce",
    "PESOMERCE": "peso_merce",
    "VALOREMERCE": "valore_merce",
    "DATAINTERVENTO": "data_intervento",
    "ALLEGATI": "allegati",
    # Tags that will be replaced by multi-paragraph content later
    "DINAMICA_EVENTI": "{{DINAMICA_EVENTI}}",
    "ACCERTAMENTI": "{{ACCERTAMENTI}}",
    "QUANTIFICAZIONE": "{{QUANTIFICAZIONE}}",
    "COMMENTO": "{{COMMENTO}}",
}

# Maps the section placeholder tags found in the template (keys)
# to the expected keys in the context dictionary that hold their content (values).
SECTION_PLACEHOLDER_TO_CONTEXT_KEY_MAPPING: dict[str, str] = {
    "{{DINAMICA_EVENTI}}": "dinamica_eventi",
    "{{ACCERTAMENTI}}": "accertamenti",
    "{{QUANTIFICAZIONE}}": "quantificazione",
    "{{COMMENTO}}": "commento",
}


def _replace_paragraph(p: ParagraphProtocol, content_string: str) -> None:
    """Replace placeholder paragraph with plain multi-paragraph content."""
    style = p.style
    p.clear()

    # Nothing to write → exit
    if not content_string or not content_string.strip():
        return

    # Split on blank lines → individual paragraphs
    blocks = [t.strip() for t in str(content_string).split("\n\n") if t.strip()]
    if not blocks:
        return

    # First paragraph re-uses the cleared placeholder
    p.add_run(blocks[0])

    # Remaining paragraphs follow directly after
    current_p = p
    for block in blocks[1:]:
        new_para = current_p.insert_paragraph_after(text="", style=style)
        new_para.add_run(block)
        current_p = new_para  # ← keeps the chain contiguous


async def inject(template_path: str, context: ReportContext) -> bytes:
    """Inject ReportContext content into the document template asynchronously."""

    def _sync_doc_generation(template_path_str: str, report_context: ReportContext) -> bytes:
        request_id = str(uuid4())
        logger.info("[%s] Starting document generation with template: %s", request_id, template_path_str)
        try:
            tpl = DocxTemplate(template_path_str)
            logger.debug("[%s] Successfully loaded template", request_id)
            mapping_data = {
                tpl_tag: getattr(report_context, ctx_key, None)
                for tpl_tag, ctx_key in TEMPLATE_TAG_TO_CONTEXT_KEY_MAPPING.items()
                if (ctx_key not in SECTION_PLACEHOLDER_TO_CONTEXT_KEY_MAPPING.values() and tpl_tag not in SECTION_PLACEHOLDER_TO_CONTEXT_KEY_MAPPING)
            }

            # Add section placeholders to context
            for _tag, _key in SECTION_PLACEHOLDER_TO_CONTEXT_KEY_MAPPING.items():
                mapping_data[_tag.removeprefix("{{").removesuffix("}}")] = _tag

            # Render the template with context data
            tpl.render(mapping_data)
            logger.debug("[%s] Successfully rendered template with mapping_data", request_id)

            # Convert to Document object for further processing
            bio = io.BytesIO()
            tpl.save(bio)
            bio.seek(0)
            doc = Document(bio)

            # Create a map of section content
            section_content_map = {placeholder: getattr(report_context, ctx_key, "") for placeholder, ctx_key in SECTION_PLACEHOLDER_TO_CONTEXT_KEY_MAPPING.items()}

            # Create a set to track consumed placeholders (for validation)
            consumed_placeholders = set()

            # Freeze the paragraph list before iterating to prevent re-visiting newly inserted paragraphs
            paragraphs = list(doc.paragraphs)

            # Build a reverse lookup for efficient placeholder detection
            placeholder_lookup = {}
            for tag in SECTION_PLACEHOLDER_TO_CONTEXT_KEY_MAPPING.keys():
                placeholder_lookup[tag] = tag

            # Process each paragraph in the frozen list
            for p_raw in paragraphs:
                p = cast(ParagraphProtocol, p_raw)  # Cast to our Protocol type
                current_text = p.text

                # Check for placeholders more efficiently
                for tag in placeholder_lookup:
                    if tag in current_text:
                        content_string = section_content_map.get(tag, "")
                        _replace_paragraph(p, content_string)
                        consumed_placeholders.add(tag)
                        logger.debug(
                            "[%s] Processed section for placeholder: %s",
                            request_id,
                            tag,
                        )
                        break  # Early exit after filling a placeholder

            # Check if all placeholders were consumed
            missing_placeholders = set(section_content_map.keys()) - consumed_placeholders
            if missing_placeholders:
                logger.warning(
                    "[%s] Some section placeholders were not found in the template: %s",
                    request_id,
                    ", ".join(missing_placeholders),
                )

            # Save the document to bytes
            out = io.BytesIO()
            doc.save(out)
            out.seek(0)
            result = out.read()
            logger.info(
                "[%s] Successfully generated document, size: %d bytes",
                request_id,
                len(result),
            )
            return result
        except DocBuilderError:
            raise
        except Exception as e:
            logger.exception("[%s] Unexpected error in document generation", request_id)
            raise DocBuilderError("Unexpected error in document generation") from e

    return await asyncio.to_thread(_sync_doc_generation, template_path, context)
