import io
import json
from pathlib import Path

import pytest
from docx import Document

from app.services.doc_builder import DocBuilderError, _add_markdown, inject

FIXTURES = Path(__file__).parent.parent / "fixtures"


def test_add_markdown_basic():
    doc = Document()
    par = doc.add_paragraph()
    _add_markdown(par, "Hello **World**!")
    runs = par.runs
    assert len(runs) == 3
    assert runs[0].text == "Hello "
    assert runs[1].text == "World"
    assert runs[1].bold is True
    assert runs[2].text == "!"


def test_inject_valid_payload():
    tpl_path = "app/templates/template.docx"
    # Load JSON payload from fixtures/payload.json
    payload_path = FIXTURES / "payload.json"
    payload_str = payload_path.read_text()
    result = inject(tpl_path, payload_str)
    assert isinstance(result, bytes)
    assert result[:2] == b"PK"
    # Verify placeholders have been replaced (no section tags remain)
    doc = Document(io.BytesIO(result))
    texts = [p.text for p in doc.paragraphs]
    for tag in (
        "{{DINAMICA_EVENTI}}",
        "{{ACCERTAMENTI}}",
        "{{QUANTIFICAZIONE}}",
        "{{COMMENTO}}",
    ):
        assert not any(tag in t for t in texts)


def test_inject_invalid_json():
    tpl_path = "app/templates/template.docx"
    with pytest.raises(DocBuilderError) as exc:
        inject(tpl_path, "not a json")
    assert "Invalid JSON payload" in str(exc.value)


def test_inject_missing_field():
    tpl_path = "app/templates/template.docx"
    payload = json.dumps({})
    with pytest.raises(DocBuilderError) as exc:
        inject(tpl_path, payload)
    assert "Missing required field" in str(exc.value)


def test_inject_invalid_template():
    tpl_path = "nonexistent.docx"
    minimal_payload = {
        "client": "",
        "client_address1": "",
        "client_address2": "",
        "date": "",
        "vs_rif": "",
        "rif_broker": "",
        "polizza": "",
        "ns_rif": "",
        "assicurato": "",
        "indirizzo_ass1": "",
        "indirizzo_ass2": "",
        "luogo": "",
        "data_danno": "",
        "cause": "",
        "data_incarico": "",
        "merce": "",
        "peso_merce": "",
        "valore_merce": "",
        "data_intervento": "",
        "dinamica_eventi": "",
        "accertamenti": "",
        "quantificazione": "",
        "commento": "",
        "allegati": "",
    }
    with pytest.raises(DocBuilderError) as exc:
        inject(tpl_path, json.dumps(minimal_payload))
    assert "Failed to render template" in str(exc.value)
